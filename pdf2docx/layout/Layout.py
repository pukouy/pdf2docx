# -*- coding: utf-8 -*-

'''
Layout objects based on PDF raw dict extracted with PyMuPDF.

@created: 2020-07-22
@author: train8808@gmail.com
---

The raw page content extracted with PyMuPDF, `page.getText('rawdict')` is described per link:
https://pymupdf.readthedocs.io/en/latest/textpage.html

In addition to the raw layout dict, some new features are also included, e.g.
    - page margin
    - rectangle shapes, for text format, annotations and table border/shading
    - new block in table type

{
    # raw dict
    ----------------------------
    "width" : w,
    "height": h,    
    "blocks": [{...}, {...}, ...],

    # introduced dict
    ----------------------------
    "margin": [left, right, top, bottom],
    "rects" : [{...}, {...}, ...]
}
'''


import json
from docx.shared import Pt
from docx.enum.section import WD_SECTION

from .Blocks import Blocks
from ..shape.Rectangles import Rectangles
from ..table.TablesConstructor import TablesConstructor
from ..common.base import PlotControl
from ..common.utils import debug_plot, DM, ITP
from ..common.pdf import new_page_with_margin
from ..common.docx import reset_paragraph_format


class Layout:
    ''' Object representing the whole page, e.g. margins, blocks, shapes, spacing.'''

    def __init__(self, raw:dict) -> None:
        self.width = raw.get('width', 0.0)
        self.height = raw.get('height', 0.0)

        # initialize blocks
        self.blocks = Blocks().from_dicts(raw.get('blocks', []))

        # initialize rects: to add rectangles later
        self.rects = Rectangles()

        # table parser
        self._tables_constructor = TablesConstructor(self.blocks, self.rects)

        # page margin: to calculate after cleaning blocks
        self._margin = None

    @property
    def margin(self):
        return self._margin

    
    @property
    def bbox_raw(self):
        if self._margin is None:
            return (0,) * 4
        else:
            left, right, top, bottom = self.margin
            return (left, top, self.width-right, self.height-bottom)


    def store(self):
        return {
            'width': self.width,
            'height': self.height,
            'margin': self._margin,
            'blocks': self.blocks.store(),
            'rects': self.rects.store(),
        }


    def serialize(self, filename:str):
        '''Write layout to specified file.'''
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(json.dumps(self.store(), indent=4))


    def plot(self, doc, title:str, key:PlotControl=PlotControl.LAYOUT):
        '''Plot specified type of blocks layout with PyMuPDF.
            ---
            Args:
              - doc: fitz.Document object
        '''
        # get objects to plot
        #  - all blocks
        if key == PlotControl.LAYOUT: 
            objects = list(self.blocks)
        
        #  - explicit table structure only
        elif key == PlotControl.TABLE: 
            objects = self.blocks.explicit_table_blocks
        
        #  - implicit table structure only
        elif key == PlotControl.IMPLICIT_TABLE: 
            objects = self.blocks.implicit_table_blocks
        
        #  - rectangle shapes
        elif key == PlotControl.SHAPE: 
            objects = list(self.rects)

        else:
            objects = []

        # do nothing if no objects
        if not objects: return

        # insert a new page
        page = new_page_with_margin(doc, self.width, self.height, self.margin, title)

        # plot styled table but no text blocks in cell
        if key==PlotControl.TABLE: 
            for item in objects:
                item.plot(page, content=False, style=True)
        
        # plot non-styled table and no text blocks in cell
        elif key==PlotControl.IMPLICIT_TABLE: 
            for item in objects:
                item.plot(page, content=False, style=False)
        
        else:
            for item in objects:
                 item.plot(page) # default args for TableBlock.plot


    def parse(self, **kwargs):
        ''' Parse page layout.
            ---
            Args:
              - kwargs: dict for layout plotting
                    kwargs = {
                        'debug': bool,
                        'doc': fitz.Document object or None,
                        'filename': str
                    }
        '''

        # preprocessing, e.g. change block order, clean negative block
        self.clean(**kwargs)        
    
        # parse table blocks: 
        #  - table structure/format recognized from rectangles
        self.parse_explicit_tables(**kwargs)
        
        #  - cell contents extracted from text blocks
        self.parse_implicit_tables(**kwargs)

        # parse text format, e.g. highlight, underline
        self.parse_text_format(**kwargs)
        
        # paragraph / line spacing
        self.parse_vertical_spacing()


    def extract_tables(self):
        '''Extract content from explicit tables.'''
        # parsing tables
        self.clean().parse_explicit_tables()

        # check table
        tables = [] # type: list[ list[list[str]] ]
        for table_block in self.blocks.table_blocks:
            tables.append(table_block.text)

        return tables


    def make_page(self, doc):
        ''' Create page based on layout data. 

            To avoid incorrect page break from original document, a new page section
            is created for each page.

            Support general document style only:
              - writing mode: from left to right, top to bottom
              - text direction: horizontal

            The vertical postion of paragraph/table is defined by space_before or 
            space_after property of a paragraph.
        '''
        # new page section
        # a default section is created when initialize the document,
        # so we do not have to add section for the first time.
        if not doc.paragraphs:
            section = doc.sections[0]
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)

        section.page_width  = Pt(self.width)
        section.page_height = Pt(self.height)

        # set page margin
        left,right,top,bottom = self.margin
        section.left_margin = Pt(left)
        section.right_margin = Pt(right)
        section.top_margin = Pt(top)
        section.bottom_margin = Pt(bottom)

        # add paragraph or table according to parsed block
        for block in self.blocks:           

            # make paragraphs
            if block.is_text_block():
                # new paragraph
                p = doc.add_paragraph()
                block.make_docx(p, self.bbox_raw)
            
            # make table
            elif block.is_table_block():

                # create dummy paragraph if table before space is set
                # line spacing: table before_space/2.0
                # before space: table before space / 2.0
                if block.before_space:
                    h = round(block.before_space/2.0, 1)
                    p = doc.add_paragraph()
                    pf = reset_paragraph_format(p)
                    pf.space_before = Pt(max(h, 0.0))
                    pf.space_after = Pt(0)
                    pf.line_spacing = Pt(h)

                # new table            
                table = doc.add_table(rows=block.num_rows, cols=block.num_cols)
                table.autofit = False
                table.allow_autofit  = False
                block.make_docx(table, self.margin)
                
        # NOTE: If a table is at the end of a page, a new paragraph will be automatically 
        # added by the rending engine, e.g. MS Word, which resulting in an unexpected
        # page break. The solution is to never put a table at the end of a page, so add
        # an empty paragraph and reset its format, particularly line spacing, when a table
        # is created.
        if len(self.blocks) and self.blocks[-1].is_table_block():
            p = doc.add_paragraph()
            reset_paragraph_format(p, Pt(1.0)) # a small line height: 1 Pt


    @debug_plot('Clean Blocks and Shapes', plot=True, category=PlotControl.SHAPE)
    def clean(self, **kwargs):
        '''Clean blocks and rectangles, e.g. remove negative blocks, duplicated rects.'''
        page_bbox = (0.0, 0.0, self.width, self.height)
        clean_blocks = self.blocks.clean(page_bbox)
        clean_rects  = self.rects.clean(page_bbox)
        
        # calculate page margin based on clean layout
        self._margin = self.page_margin()

        return clean_blocks or clean_rects


    @debug_plot('Explicit Table Structure', plot=True, category=PlotControl.TABLE)
    def parse_explicit_tables(self, **kwargs) -> bool:
        '''parse table structure from rectangle shapes'''
        tables = self._tables_constructor.explicit_tables()
        return bool(tables)


    @debug_plot('Implicit Table Structure', plot=True, category=PlotControl.IMPLICIT_TABLE)
    def parse_implicit_tables(self, **kwargs):
        ''' Parse table structure based on the layout of text/image blocks.

            Since no cell borders exist in this case, there may be various probabilities of table structures. 
            Among which, we use the simplest one, i.e. 1-row and n-column, to make the docx look like pdf.

            Ensure no horizontally aligned blocks in each column, so that these blocks can be converted to
            paragraphs consequently in docx.
        '''
        # horizontal range of table
        left, right, *_ = self.margin
        X0, X1 = left, self.width - right

        tables = self._tables_constructor.implicit_tables(X0, X1)
        return bool(tables)


    @debug_plot('Parsed Text Blocks', plot=True, category=PlotControl.LAYOUT)
    def parse_text_format(self, **kwargs):
        '''Parse text format in both page and table context.'''
        return self.blocks.parse_text_format(self.rects)
 

    def page_margin(self):
        '''Calculate page margin.            
            ---
            Args:
            - width: page width
            - height: page height

            Calculation method:
            - left: MIN(bbox[0])
            - right: MIN(left, width-max(bbox[2]))
            - top: MIN(bbox[1])
            - bottom: height-MAX(bbox[3])
        '''
        # return normal page margin if no blocks exist
        if not self.blocks and not self.rects:
            return (ITP, ) * 4                 # 1 Inch = 72 pt

        # consider both blocks and rects for page margin
        list_bbox = list(map(lambda x: x.bbox, self.blocks))
        list_bbox.extend(list(map(lambda x: x.bbox, self.rects))) 

        # left margin 
        left = min(map(lambda x: x.x0, list_bbox))
        left = max(left, 0)

        # right margin
        x_max = max(map(lambda x: x.x1, list_bbox))
        right = self.width - x_max - DM*10.0  # consider tolerance: leave more free space
        right = min(right, left)              # symmetry margin if necessary
        right = max(right, 0.0)               # avoid negative margin

        # top margin
        top = min(map(lambda x: x.y0, list_bbox))
        top = max(top, 0)

        # bottom margin
        bottom = self.height-max(map(lambda x: x.y1, list_bbox))
        bottom = max(bottom, 0.0)

        # margin is calculated based on text block only, without considering shape, e.g. table border,
        # so reduce calculated top/bottom margin to left some free space
        top *= 0.5
        bottom *= 0.5

        # use normal margin if calculated margin is large enough
        return (
            min(ITP, left), 
            min(ITP, right), 
            min(ITP, top), 
            min(ITP, bottom)
            )
 

    def parse_vertical_spacing(self):
        ''' Calculate external and internal vertical space for paragraph blocks under page context 
            or table context. It'll used as paragraph spacing and line spacing when creating paragraph.
        '''
        self.blocks.parse_vertical_spacing(self.bbox_raw)
